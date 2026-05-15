"""
File format handlers — TXT, DOCX, PDF, XLSX, RTF.
All anonymization uses anonymize_text_pipeline() which returns
(anonymized_text, replacements_dict).
"""
import os
import sys
from pathlib import Path

ALLOWED_EXTENSIONS = {'.txt', '.docx', '.pdf', '.xlsx', '.rtf'}


# ─────────────────────────────────────────────────────────────────────────────
# Validation
# ─────────────────────────────────────────────────────────────────────────────

def validate_file(filepath: Path):
    ext = filepath.suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        return False, (f'Формат «{ext}» не поддерживается. '
                       f'Поддерживаются: {", ".join(sorted(ALLOWED_EXTENSIONS))}')
    if ext == '.pdf' and not _pdf_has_text_layer(filepath):
        return False, ('PDF не содержит текстового слоя (скан). '
                       'Обработка без OCR невозможна.')
    return True, ''


def _pdf_has_text_layer(filepath: Path) -> bool:
    try:
        import fitz
        doc = fitz.open(str(filepath))
        return any(page.get_text().strip() for page in doc)
    except Exception:
        return False


# ─────────────────────────────────────────────────────────────────────────────
# Entry point
# ─────────────────────────────────────────────────────────────────────────────

def process_uploaded_file(input_path: Path, output_dir: Path,
                          session_id: str, db_path, mode: str,
                          use_spacy: bool = True,
                          use_llm: bool = False) -> dict:
    valid, err = validate_file(input_path)
    if not valid:
        raise ValueError(err)

    ext    = input_path.suffix.lower()
    prefix = 'anonymized' if mode == 'anonymize' else 'deanonymized'
    stem   = input_path.stem
    for pfx in ('anonymized_', 'deanonymized_', 'anon_', 'deanon_'):
        if stem.lower().startswith(pfx):
            stem = stem[len(pfx):]
            break

    out_ext = '.docx' if ext == '.pdf' else ext
    output_name = f'{prefix}_{stem}{out_ext}'
    output_path = output_dir / output_name

    if mode == 'anonymize':
        result = _anonymize(input_path, output_path, ext, session_id, db_path,
                            use_spacy, use_llm)
    else:
        result = _deanonymize(input_path, output_path, ext, session_id, db_path)

    result['output_filename'] = output_name
    if '_fallback_path' in result:
        fallback = Path(result.pop('_fallback_path'))
        result['output_filename'] = fallback.name
    return result


# ─────────────────────────────────────────────────────────────────────────────
# Anonymize dispatcher
# ─────────────────────────────────────────────────────────────────────────────

def _anonymize(input_path, output_path, ext, session_id, db_path,
               use_spacy=True, use_llm=False):
    if ext == '.txt':  return _anon_txt(input_path, output_path, session_id, db_path, use_spacy, use_llm)
    if ext == '.docx': return _anon_docx(input_path, output_path, session_id, db_path, use_spacy, use_llm)
    if ext == '.pdf':  return _anon_pdf(input_path, output_path, session_id, db_path, use_spacy, use_llm)
    if ext == '.xlsx': return _anon_xlsx(input_path, output_path, session_id, db_path, use_spacy, use_llm)
    if ext == '.rtf':  return _anon_rtf(input_path, output_path, session_id, db_path, use_spacy, use_llm)
    raise ValueError(f'Unsupported: {ext}')


# ─────────────────────────────────────────────────────────────────────────────
# Deanonymize dispatcher
# ─────────────────────────────────────────────────────────────────────────────

def _deanonymize(input_path, output_path, ext, session_id, db_path):
    from core.db       import get_reverse_mappings
    from core.anonymizer import apply_reverse

    rev = get_reverse_mappings(db_path, session_id)

    if ext == '.txt':
        text = input_path.read_text(encoding='utf-8', errors='replace')
        output_path.write_text(apply_reverse(text, rev), encoding='utf-8')
        return {}
    if ext == '.docx': return _deanon_docx(input_path, output_path, rev)
    if ext == '.pdf':  return _deanon_pdf(input_path, output_path, rev)
    if ext == '.xlsx': return _deanon_xlsx(input_path, output_path, rev)
    if ext == '.rtf':  return _deanon_rtf(input_path, output_path, rev)
    raise ValueError(f'Unsupported: {ext}')


# ─────────────────────────────────────────────────────────────────────────────
# DOCX helpers
# ─────────────────────────────────────────────────────────────────────────────

def _accept_tracked_changes(doc):
    from docx.oxml.ns import qn
    body = doc.element.body
    for ins in list(body.iter(qn('w:ins'))):
        parent = ins.getparent()
        if parent is None:
            continue
        idx = list(parent).index(ins)
        for child in list(ins):
            parent.insert(idx, child)
            idx += 1
        parent.remove(ins)
    for del_elem in list(body.iter(qn('w:del'))):
        parent = del_elem.getparent()
        if parent is not None:
            parent.remove(del_elem)


def _replace_para(para, replacements: dict):
    if not replacements or not para.text.strip():
        return
    sorted_reps = sorted(replacements.items(), key=lambda x: -len(x[0]))

    for old, new in sorted_reps:
        for run in para.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)

    for old, new in sorted_reps:
        if old not in para.text:
            continue
        runs = para.runs
        if not runs:
            continue
        full = ''.join(r.text for r in runs)
        if old not in full:
            continue
        new_full = full.replace(old, new)
        runs[0].text = new_full
        for r in runs[1:]:
            r.text = ''

    try:
        from docx.oxml.ns import qn
        for hl in list(para._p.iter(qn('w:hyperlink'))):
            hl_modified = False
            for run_el in hl.iter(qn('w:r')):
                for t_el in run_el.iter(qn('w:t')):
                    if t_el.text:
                        for old, new in sorted_reps:
                            if old in t_el.text:
                                t_el.text = t_el.text.replace(old, new)
                                hl_modified = True
            if hl_modified:
                parent = hl.getparent()
                if parent is not None:
                    idx = list(parent).index(hl)
                    for child in list(hl):
                        parent.insert(idx, child)
                        idx += 1
                    parent.remove(hl)
    except Exception:
        pass


def _iter_all_paras(doc):
    yield from doc.paragraphs
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in doc.sections:
        try:
            yield from section.header.paragraphs
            yield from section.footer.paragraphs
        except Exception:
            pass


def _anon_txt(input_path, output_path, session_id, db_path,
              use_spacy=True, use_llm=False):
    from core.anonymizer import anonymize_text_pipeline
    text      = input_path.read_text(encoding='utf-8', errors='replace')
    out, reps = anonymize_text_pipeline(text, db_path, session_id,
                                        use_spacy=use_spacy, use_llm=use_llm)
    output_path.write_text(out, encoding='utf-8')
    return {'entities_found': len(reps)}


def _anon_docx(input_path, output_path, session_id, db_path,
               use_spacy=True, use_llm=False):
    from docx import Document
    from core.anonymizer import anonymize_text_pipeline

    doc = Document(str(input_path))
    _accept_tracked_changes(doc)

    all_replacements = {}

    for para in _iter_all_paras(doc):
        if not para.text.strip():
            continue
        _, reps = anonymize_text_pipeline(
            para.text, db_path, session_id,
            use_spacy=use_spacy, use_llm=use_llm,
        )
        if reps:
            all_replacements.update(reps)
            _replace_para(para, reps)

    doc.save(str(output_path))
    return {'entities_found': len(all_replacements)}


def _deanon_docx(input_path, output_path, rev):
    from docx import Document

    doc = Document(str(input_path))
    _accept_tracked_changes(doc)

    if rev:
        reps = {}
        for tok, orig in rev.items():
            reps[f'[{tok}]'] = orig
            reps[tok]        = orig

        for para in _iter_all_paras(doc):
            _replace_para(para, reps)

    doc.save(str(output_path))
    return {}


# ─────────────────────────────────────────────────────────────────────────────
# PDF helpers
# ─────────────────────────────────────────────────────────────────────────────

def _detect_avg_fontsize(page) -> float:
    sizes = []
    for block in page.get_text('dict').get('blocks', []):
        for line in block.get('lines', []):
            for span in line.get('spans', []):
                sz = span.get('size', 0)
                if 6 < sz < 30:
                    sizes.append(sz)
    if not sizes:
        return 10.0
    sizes.sort()
    return sizes[len(sizes) // 2]


def _find_cyrillic_font() -> str | None:
    candidates = []
    if sys.platform == 'win32':
        candidates = [r'C:\Windows\Fonts\arial.ttf',
                      r'C:\Windows\Fonts\calibri.ttf',
                      r'C:\Windows\Fonts\times.ttf']
    elif sys.platform == 'darwin':
        candidates = ['/Library/Fonts/Arial.ttf',
                      '/System/Library/Fonts/Supplemental/Arial.ttf']
    else:
        candidates = ['/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf',
                      '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf']
    return next((p for p in candidates if os.path.exists(p)), None)


def _anon_pdf(input_path, output_path, session_id, db_path,
              use_spacy=True, use_llm=False):
    """
    Extract text from PDF via PyMuPDF, anonymize, then redact in-place.
    Output is always PDF (no pdf2docx dependency).
    For DOCX output the user can upload the original DOCX.
    """
    import fitz
    from core.anonymizer import anonymize_text_pipeline

    doc = fitz.open(str(input_path))
    all_text = '\n'.join(p.get_text() for p in doc)
    _, reps = anonymize_text_pipeline(all_text, db_path, session_id,
                                      use_spacy=use_spacy, use_llm=use_llm)

    if reps:
        sorted_reps = sorted(reps.items(), key=lambda x: -len(x[0]))
        for page in doc:
            avg_fsize = _detect_avg_fontsize(page)
            for orig, token in sorted_reps:
                for rect in page.search_for(orig):
                    page.add_redact_annot(
                        rect, text=token,
                        fontsize=max(avg_fsize * 0.88, 7),
                        fill=(1, 1, 1), text_color=(0, 0, 0),
                    )
            page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)

    pdf_out = output_path.with_suffix('.pdf')
    doc.save(str(pdf_out), garbage=4, deflate=True)
    return {'entities_found': len(reps), '_fallback_path': str(pdf_out)}


def _deanon_pdf(input_path, output_path, rev):
    import fitz
    from core.anonymizer import apply_reverse

    if not rev:
        import shutil
        shutil.copy2(str(input_path), str(output_path.with_suffix('.pdf')))
        return {'_fallback_path': str(output_path.with_suffix('.pdf'))}

    cyrillic_font = _find_cyrillic_font()
    doc = fitz.open(str(input_path))

    search_pairs = []
    for tok, orig in rev.items():
        search_pairs.append((f'[{tok}]', orig))
        search_pairs.append((tok, orig))
    search_pairs.sort(key=lambda x: -len(x[0]))

    for page in doc:
        avg_fsize = _detect_avg_fontsize(page)
        font_ok = False
        if cyrillic_font:
            try:
                page.insert_font(fontname='CyrF', fontfile=cyrillic_font)
                font_ok = True
            except Exception:
                pass

        for token_str, orig in search_pairs:
            for rect in page.search_for(token_str):
                page.draw_rect(rect, color=None, fill=(1, 1, 1))
                if font_ok:
                    try:
                        page.insert_textbox(
                            rect.inflate(1, 1), orig,
                            fontname='CyrF',
                            fontsize=max(avg_fsize * 0.88, 7),
                            color=(0, 0, 0), align=0,
                        )
                    except Exception:
                        pass

    pdf_out = output_path.with_suffix('.pdf')
    doc.save(str(pdf_out), garbage=4, deflate=True)
    return {'_fallback_path': str(pdf_out)}


# ─────────────────────────────────────────────────────────────────────────────
# XLSX helpers
# ─────────────────────────────────────────────────────────────────────────────

def _anon_xlsx(input_path, output_path, session_id, db_path,
               use_spacy=True, use_llm=False):
    from openpyxl import load_workbook
    from core.anonymizer import anonymize_text_pipeline

    wb = load_workbook(str(input_path))
    all_text = '\n'.join(
        str(cell.value) for ws in wb.worksheets
        for row in ws.iter_rows() for cell in row
        if isinstance(cell.value, str) and cell.value.strip()
    )
    _, reps = anonymize_text_pipeline(all_text, db_path, session_id,
                                      use_spacy=use_spacy, use_llm=use_llm)

    if reps:
        sorted_reps = sorted(reps.items(), key=lambda x: -len(x[0]))
        for ws in wb.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    if isinstance(cell.value, str):
                        for orig, tok in sorted_reps:
                            cell.value = cell.value.replace(orig, tok)

    wb.save(str(output_path))
    return {'entities_found': len(reps)}


def _deanon_xlsx(input_path, output_path, rev):
    from openpyxl import load_workbook
    from core.anonymizer import apply_reverse

    wb = load_workbook(str(input_path))
    if rev:
        for ws in wb.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    if isinstance(cell.value, str):
                        cell.value = apply_reverse(cell.value, rev)
    wb.save(str(output_path))
    return {}


# ─────────────────────────────────────────────────────────────────────────────
# RTF helpers
# ─────────────────────────────────────────────────────────────────────────────

def _extract_rtf_plain(filepath: Path) -> str:
    try:
        from striprtf.striprtf import rtf_to_text
        raw = filepath.read_bytes().decode('utf-8', errors='replace')
        return rtf_to_text(raw)
    except Exception:
        return ''


def _anon_rtf(input_path, output_path, session_id, db_path,
              use_spacy=True, use_llm=False):
    from core.anonymizer import anonymize_text_pipeline

    plain   = _extract_rtf_plain(input_path)
    _, reps = anonymize_text_pipeline(plain, db_path, session_id,
                                      use_spacy=use_spacy, use_llm=use_llm)
    raw        = input_path.read_bytes().decode('utf-8', errors='replace')
    for orig, tok in sorted(reps.items(), key=lambda x: -len(x[0])):
        raw = raw.replace(orig, tok)
    output_path.write_text(raw, encoding='utf-8')
    return {'entities_found': len(reps)}


def _deanon_rtf(input_path, output_path, rev):
    from core.anonymizer import apply_reverse
    raw = input_path.read_bytes().decode('utf-8', errors='replace')
    output_path.write_text(apply_reverse(raw, rev), encoding='utf-8')
    return {}
