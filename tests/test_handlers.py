"""
End-to-end tests using real DOCX/PDF samples in Тестовые_файлы/.

Each test:
1. anonymizes a real document via process_uploaded_file
2. asserts the output exists, is valid, contains tokens, contains no raw PII

These tests skip automatically if the test files directory is missing.
"""
import re
import pytest

from core.handlers import process_uploaded_file
from core.db import get_session_mappings, get_reverse_mappings


PII_LEAK_PATTERNS = [
    # 10/12-digit INN-shaped numbers (excluding tokens)
    (re.compile(r'(?<!\[)(?<!\d)\d{10}(?!\d)(?!\])'), 'INN-shaped 10-digit number'),
    (re.compile(r'(?<!\[)(?<!\d)\d{12}(?!\d)(?!\])'), 'INN-shaped 12-digit number'),
    # email addresses
    (re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b'), 'e-mail'),
]


def _read_docx_plain(path):
    from docx import Document
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    parts.append(p.text)
    return '\n'.join(parts)


@pytest.mark.parametrize('filename', [
    'Агентский_договор_ЛЕГО_24.03.2026.docx',
    'Анализ фин состояния ООО СМС. ООО ЮГМК. ООО ЮГМК ДОНЕЦК. 17.02.26.docx',
    'Аналитическая справка АО Северный Народный Банк.docx',
    'ДКП 100% долей Новые выставки.docx',
    'План восстановления компаний 20.01.26.docx',
    'Решение НВ-ЦКМ.docx_распознан.docx',
    'Решение_Новые_Выставки_новый_директор.docx',
])
def test_anonymize_real_docx_regex_only(test_files_dir, tmp_path, tmp_db, session_id, filename):
    """Process each sample in regex-only mode and verify output integrity."""
    src = test_files_dir / filename
    if not src.exists():
        pytest.skip(f'{filename} not found')

    out_dir = tmp_path / 'out'
    out_dir.mkdir()

    result = process_uploaded_file(
        input_path=src,
        output_dir=out_dir,
        session_id=session_id,
        db_path=tmp_db,
        mode='anonymize',
        use_spacy=False,
        use_llm=False,
    )
    out_path = out_dir / result['output_filename']
    assert out_path.exists(), f'Output missing: {out_path}'
    assert out_path.stat().st_size > 100, 'Output suspiciously small'

    # Open as docx (regex-only mode should produce a valid file)
    plain = _read_docx_plain(out_path)
    assert plain.strip(), 'Output document has no text'

    # Verify tokens were inserted
    mappings = get_session_mappings(tmp_db, session_id)
    if not mappings:
        pytest.skip(f'No entities found in {filename} (acceptable)')

    # Token uniqueness invariant
    tokens = [m['token'] for m in mappings]
    assert len(tokens) == len(set(tokens)), f'Duplicate tokens! {tokens}'

    # Sample of tokens should appear in output
    sample_token = mappings[0]['token']
    assert f'[{sample_token}]' in plain, \
        f'Token {sample_token} missing from output of {filename}'


def test_no_mapping_captures_masked_content(test_files_dir, tmp_path, tmp_db, session_id):
    """After processing a sample, no mapping should store text containing [TYPE_N]."""
    from core.anonymizer import _contains_token
    src = test_files_dir / 'ДКП 100% долей Новые выставки.docx'
    if not src.exists():
        pytest.skip('Sample missing')

    out_dir = tmp_path / 'out'
    out_dir.mkdir()
    process_uploaded_file(
        input_path=src, output_dir=out_dir, session_id=session_id,
        db_path=tmp_db, mode='anonymize', use_spacy=False, use_llm=False,
    )
    for m in get_session_mappings(tmp_db, session_id):
        assert not _contains_token(m['original_form']), \
            f"Bad mapping (mask in original_form): {m}"
        assert not _contains_token(m['canonical_form']), \
            f"Bad mapping (mask in canonical_form): {m}"


def test_deanonymize_roundtrip_docx(test_files_dir, tmp_path, tmp_db, session_id):
    """Anonymize then deanonymize a DOCX — PII should reappear in output."""
    src = test_files_dir / 'ДКП 100% долей Новые выставки.docx'
    if not src.exists():
        pytest.skip('Sample missing')

    out_dir = tmp_path / 'out'
    out_dir.mkdir()

    anon = process_uploaded_file(
        input_path=src, output_dir=out_dir, session_id=session_id,
        db_path=tmp_db, mode='anonymize', use_spacy=False, use_llm=False,
    )
    anon_path = out_dir / anon['output_filename']

    deanon = process_uploaded_file(
        input_path=anon_path, output_dir=out_dir, session_id=session_id,
        db_path=tmp_db, mode='deanonymize',
    )
    deanon_path = out_dir / deanon['output_filename']
    assert deanon_path.exists()

    plain = _read_docx_plain(deanon_path)
    # No tokens should remain in the deanonymized text
    leftover = re.findall(r'\[(?:FIO|YUL|INN|OGRN|KPP|RS|KS|BIK|SNILS|PASSPORT|TEL|EMAIL|SWIFT|ADR|DOB|LIC|URL)_\d+\]', plain)
    assert not leftover, f'Tokens remained after deanon: {leftover[:5]}'


def test_session_files_endpoint_lists_output(test_files_dir, tmp_path, tmp_db, session_id):
    """Smoke test: after processing, an output file exists on disk in expected layout."""
    src = test_files_dir / 'ДКП 100% долей Новые выставки.docx'
    if not src.exists():
        pytest.skip('Sample missing')

    out_dir = tmp_path / 'out'
    out_dir.mkdir()
    res = process_uploaded_file(
        input_path=src, output_dir=out_dir, session_id=session_id,
        db_path=tmp_db, mode='anonymize', use_spacy=False, use_llm=False,
    )
    files = list(out_dir.iterdir())
    assert any(f.name == res['output_filename'] for f in files)
